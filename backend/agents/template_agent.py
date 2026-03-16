"""
Template Agent — 7-step smart contract generation pipeline.

Steps:
  1. Detect permit/contract type from reference document
  2. Select matching DOCX template from GCS template library
  3. Extract template layout (headers, footers, placeholders)
  4. Extract variable data from reference document (fields that change per customer)
  5. Populate DOCX template with extracted + lead data
  6. Preserve original template structure (headers, footers, page formatting)
  7. Return final DOCX bytes ready for download or GCS storage
"""

import asyncio
import io
import json
import os
from functools import partial

import pdfplumber
from docx import Document
from docx.oxml.ns import qn
from dotenv import load_dotenv
from openai import AsyncOpenAI

from db.storage_client import download_file, upload_file

load_dotenv()

_client: AsyncOpenAI | None = None

# ── Supported permit types → GCS template file names ─────────────────────── #
PERMIT_TYPE_MAP: dict[str, str] = {
    "adu_legalization": "adu_legalization.docx",
    "adu_permit": "adu_legalization.docx",
    "remodeling": "remodeling.docx",
    "kitchen_remodel": "remodeling.docx",
    "bathroom_remodel": "remodeling.docx",
    "home_addition": "home_addition.docx",
    "room_addition": "home_addition.docx",
    "electrical_permit": "electrical_permit.docx",
    "roof_replacement": "remodeling.docx",
    "deck_build": "remodeling.docx",
    "general_construction": "remodeling.docx",
}

# Variables that appear in every DOCX template as {{variable_name}} placeholders
STANDARD_VARIABLES = [
    "customer_name",
    "customer_email",
    "customer_phone",
    "customer_address",
    "city",
    "state",
    "zip_code",
    "project_type",
    "project_description",
    "scope_of_work",
    "quote_amount",
    "deposit_amount",
    "milestone_2_amount",
    "milestone_2_name",
    "final_payment_amount",
    "start_date",
    "estimated_completion",
    "contract_date",
    "company_name",
    "company_address",
    "company_phone",
    "company_license",
    "permit_number",
    "parcel_number",
]


def _get_openai() -> AsyncOpenAI:
    global _client
    if _client is None:
        _client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    return _client


# ── Step 1: Detect permit type ─────────────────────────────────────────────── #

async def detect_permit_type(document_text: str) -> str:
    """
    Use GPT-4o to identify the permit/contract type from document text.
    Returns one of the keys in PERMIT_TYPE_MAP.
    """
    result = await _get_openai().chat.completions.create(
        model="gpt-4o",
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a construction document classifier. "
                    "Identify the contract/permit type from the document text. "
                    "Return JSON with a single key 'permit_type' whose value is exactly one of: "
                    + ", ".join(PERMIT_TYPE_MAP.keys())
                    + ". If unsure, use 'general_construction'."
                ),
            },
            {"role": "user", "content": document_text[:4000]},
        ],
    )
    data = json.loads(result.choices[0].message.content or "{}")
    permit_type = data.get("permit_type", "general_construction")
    return permit_type if permit_type in PERMIT_TYPE_MAP else "general_construction"


# ── Step 2: Select template ────────────────────────────────────────────────── #

def select_template_name(permit_type: str) -> str:
    """Return the GCS template filename for the given permit type."""
    return PERMIT_TYPE_MAP.get(permit_type, "remodeling.docx")


# ── Step 3 + 4: Extract variables from reference document ─────────────────── #

async def extract_contract_variables(
    document_text: str,
    permit_type: str,
    lead_data: dict,
) -> dict:
    """
    Use GPT-4o to extract all variable fields from the reference document.
    Merges with lead_data so known fields are never missed.
    """
    result = await _get_openai().chat.completions.create(
        model="gpt-4o",
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": (
                    f"You are extracting variable fields from a {permit_type} contract document. "
                    "Extract all customer-specific and project-specific data that would change "
                    "for a different customer. Return a flat JSON object with these keys where found: "
                    + ", ".join(STANDARD_VARIABLES)
                    + ". For any field not found in the document, return an empty string. "
                    "Format amounts as numbers only (no $ signs). Format dates as YYYY-MM-DD."
                ),
            },
            {
                "role": "user",
                "content": f"Reference document:\n{document_text[:6000]}",
            },
        ],
    )
    extracted = json.loads(result.choices[0].message.content or "{}")

    # Lead data always overrides extracted data
    merged = {**extracted}
    if lead_data.get("customer_name"):
        merged["customer_name"] = lead_data["customer_name"]
    if lead_data.get("city"):
        merged["city"] = lead_data["city"]
    if lead_data.get("quote_amount"):
        merged["quote_amount"] = str(lead_data["quote_amount"])
        # Auto-calculate standard splits if not extracted: 30/40/30
        amount = float(lead_data["quote_amount"])
        if not merged.get("deposit_amount"):
            merged["deposit_amount"] = f"{amount * 0.30:,.2f}"
        if not merged.get("milestone_2_amount"):
            merged["milestone_2_amount"] = f"{amount * 0.40:,.2f}"
        if not merged.get("final_payment_amount"):
            merged["final_payment_amount"] = f"{amount * 0.30:,.2f}"
    if lead_data.get("email"):
        merged["customer_email"] = lead_data["email"]
    if lead_data.get("phone"):
        merged["customer_phone"] = lead_data["phone"]
    if lead_data.get("project_type"):
        merged["project_type"] = lead_data["project_type"]
    if lead_data.get("notes"):
        merged["project_description"] = lead_data["notes"]

    # Defaults for company fields — pulled from env so each deployment can customize
    merged.setdefault("company_name", os.getenv("COMPANY_NAME", "QTC Construction LLC"))
    merged.setdefault("company_address", os.getenv("COMPANY_ADDRESS", ""))
    merged.setdefault("company_phone", os.getenv("COMPANY_PHONE", ""))
    merged.setdefault("company_email", os.getenv("SMTP_USER", ""))
    merged.setdefault("company_license", os.getenv("COMPANY_LICENSE", ""))

    return merged


# ── Step 5 + 6: Populate DOCX template ────────────────────────────────────── #

def _replace_in_paragraph(paragraph, variables: dict) -> None:
    """Replace {{key}} placeholders in a single paragraph, preserving runs."""
    for key, value in variables.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in paragraph.text:
            # Rebuild the paragraph text across all runs
            full_text = "".join(r.text for r in paragraph.runs)
            if placeholder in full_text:
                new_text = full_text.replace(placeholder, str(value))
                # Put all text in first run, clear rest
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    for run in paragraph.runs[1:]:
                        run.text = ""


def _replace_in_element(element, variables: dict) -> None:
    """Recursively replace placeholders in paragraphs and table cells."""
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    for child in element.iter_inner_content() if hasattr(element, 'iter_inner_content') else []:
        if isinstance(child, Paragraph):
            _replace_in_paragraph(child, variables)
        elif isinstance(child, Table):
            for row in child.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_paragraph(para, variables)


def populate_docx_template(template_bytes: bytes, variables: dict) -> bytes:
    """
    Open a DOCX template and replace all {{variable}} placeholders.
    Handles: body paragraphs, tables, headers, footers.
    Returns populated DOCX as bytes.
    """
    doc = Document(io.BytesIO(template_bytes))

    # Body paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, variables)

    # Tables in body
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, variables)

    # Headers and footers in all sections
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    _replace_in_paragraph(para, variables)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                _replace_in_paragraph(para, variables)

        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    _replace_in_paragraph(para, variables)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                _replace_in_paragraph(para, variables)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


# ── Document text extraction ───────────────────────────────────────────────── #

def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from DOCX bytes."""
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from PDF bytes using pdfplumber."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
    return "\n".join(text_parts)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Detect file type and extract text."""
    fname = filename.lower()
    if fname.endswith(".pdf"):
        return _extract_text_from_pdf(file_bytes)
    elif fname.endswith(".docx"):
        return _extract_text_from_docx(file_bytes)
    else:
        # Try as plain text
        return file_bytes.decode("utf-8", errors="ignore")


# ── Step 7: Full pipeline ──────────────────────────────────────────────────── #

async def generate_from_reference(
    reference_bytes: bytes,
    filename: str,
    lead_data: dict,
    org_id: str,
) -> tuple[bytes, str, str]:
    """
    Full 7-step pipeline.

    Returns:
        (docx_bytes, contract_id, permit_type)
        docx_bytes: populated DOCX ready for download
        contract_id: Firestore contract document ID
        permit_type: detected type string
    """
    loop = asyncio.get_running_loop()

    # Step 1 — detect permit type
    ref_text = await loop.run_in_executor(None, extract_text, reference_bytes, filename)
    permit_type = await detect_permit_type(ref_text)

    # Step 2 — select template
    template_filename = select_template_name(permit_type)
    template_path = f"templates/{template_filename}"

    # Step 3 — download DOCX template from GCS
    try:
        template_bytes = await download_file(template_path)
    except Exception as e:
        # Fallback: generate minimal DOCX if GCS template not found
        import logging
        logging.warning(f"GCS template '{template_path}' not found ({e}), using fallback")
        template_bytes = await loop.run_in_executor(
            None, _create_fallback_docx_template, permit_type
        )

    # Step 4 — extract variables from reference document
    variables = await extract_contract_variables(ref_text, permit_type, lead_data)

    # Step 5 + 6 — populate DOCX template preserving layout
    populated_bytes = await loop.run_in_executor(
        None, populate_docx_template, template_bytes, variables
    )

    # Step 7 — upload populated contract to GCS and save to Firestore
    import uuid
    from google.cloud.firestore import SERVER_TIMESTAMP
    from db.firestore_client import get_db

    contract_id = str(uuid.uuid4())
    storage_path = f"contracts/{org_id}/{contract_id}.docx"
    await upload_file(populated_bytes, storage_path)

    db = get_db()
    await db.collection("contracts").document(contract_id).set({
        "lead_id": lead_data.get("id", ""),
        "template_name": permit_type,
        "content": ref_text[:500],  # preview only
        "status": "draft",
        "signed_at": None,
        "storage_path": storage_path,
        "org_id": org_id,
        "created_at": SERVER_TIMESTAMP,
        "variables": variables,
    })

    return populated_bytes, contract_id, permit_type


# ── Fallback: create a minimal DOCX if no template exists in GCS ──────────── #

def _create_fallback_docx_template(permit_type: str) -> bytes:
    """
    Create a minimal DOCX template with standard placeholders.
    Used when no DOCX template exists in GCS for the given permit type.
    """
    doc = Document()

    # Header
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "QTC Construction LLC  |  {{company_address}}  |  {{company_phone}}"

    # Title
    title = doc.add_heading(permit_type.replace("_", " ").title() + " Contract", 0)

    doc.add_paragraph(f"Contract Date: {{{{contract_date}}}}")
    doc.add_paragraph(f"Permit #: {{{{permit_number}}}}")
    doc.add_paragraph("")

    doc.add_heading("Parties", level=1)
    doc.add_paragraph(
        "This agreement is between QTC Construction LLC (\"Contractor\") "
        "and {{customer_name}} (\"Client\") located at {{customer_address}}, "
        "{{city}}, {{state}} {{zip_code}}."
    )
    doc.add_paragraph("Client Phone: {{customer_phone}}  |  Email: {{customer_email}}")

    doc.add_heading("Scope of Work", level=1)
    doc.add_paragraph("{{scope_of_work}}")
    doc.add_paragraph("Project Description: {{project_description}}")

    doc.add_heading("Payment Schedule", level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Milestone"
    hdr[1].text = "Amount"
    hdr[2].text = "Due"
    row1 = table.rows[1].cells
    row1[0].text = "Deposit (30%)"
    row1[1].text = "${{deposit_amount}}"
    row1[2].text = "Upon signing"
    row2 = table.rows[2].cells
    row2[0].text = "{{milestone_2_name}}"
    row2[1].text = "${{milestone_2_amount}}"
    row2[2].text = "Mid-project"
    row3 = table.rows[3].cells
    row3[0].text = "Final Payment (30%)"
    row3[1].text = "${{final_payment_amount}}"
    row3[2].text = "Upon completion"

    doc.add_paragraph("")
    doc.add_paragraph(f"Total Contract Amount: ${{{{quote_amount}}}}")

    doc.add_heading("Timeline", level=1)
    doc.add_paragraph("Estimated Start: {{start_date}}")
    doc.add_paragraph("Estimated Completion: {{estimated_completion}}")

    doc.add_heading("Signatures", level=1)
    doc.add_paragraph("Client: _______________________________  Date: __________")
    doc.add_paragraph("Contractor: ____________________________  Date: __________")
    doc.add_paragraph("License #: {{company_license}}")

    # Footer
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "QTC Construction LLC  |  Confidential  |  Page 1"

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
